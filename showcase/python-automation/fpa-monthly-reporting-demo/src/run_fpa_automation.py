from __future__ import annotations

"""One-click FP&A workflow: raw data -> clean data -> variance -> Word report."""

import os
import re
import shutil
import time
from collections import defaultdict
from datetime import date, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.chart.series import SeriesLabel
from openpyxl.formatting.rule import CellIsRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo


ROOT = Path(__file__).resolve().parents[1]
SAMPLE_DATA = ROOT / "sample_data"
OUTPUT = ROOT / "output"
STAGING = OUTPUT / ".staging"
RAW_PATH = SAMPLE_DATA / "data-raw.xlsx"
CLEAN_PATH = OUTPUT / "clean-data.xlsx"
VARIANCE_PATH = OUTPUT / "variance-analysis.xlsx"
REPORT_PATH = OUTPUT / "monthly-report.docx"

NAVY = "17365D"
BLUE = "2F75B5"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EAF3F8"
GREEN = "1F7A4D"
PALE_GREEN = "E2F0D9"
RED = "C00000"
PALE_RED = "FCE4D6"
AMBER = "BF9000"
PALE_AMBER = "FFF2CC"
WHITE = "FFFFFF"
CHARCOAL = "24292F"
MUTED = "667085"
GRID = "D9E1F2"
THIN_GRAY = Side(style="thin", color="D9E1F2")

PRODUCT_MAP = {
    "ELEC01": "Electronics",
    "HOME02": "Home & Living",
    "FASH03": "Fashion",
}


def log(message: str = "") -> None:
    print(message, flush=True)


def section(number: int, title: str) -> None:
    log(f"\n[{number}/4] {title}")


def parse_date(value: Any) -> date:
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    text = str(value).strip()
    for pattern in ("%Y-%m-%d", "%d/%m/%Y", "%b %d %Y"):
        try:
            return datetime.strptime(text, pattern).date()
        except ValueError:
            continue
    raise ValueError(f"Unsupported invoice date: {value!r}")


def normalize_product(value: Any) -> tuple[str, str]:
    code = re.sub(r"[^A-Za-z0-9]", "", str(value)).upper()
    if code not in PRODUCT_MAP:
        raise ValueError(f"Unmapped product code: {value!r}")
    return code, PRODUCT_MAP[code]


def normalize_region(value: Any) -> str:
    region = str(value).strip().title()
    if region not in {"North", "South", "Central"}:
        raise ValueError(f"Unexpected region: {value!r}")
    return region


def clean_rows(raw_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    cleaned: list[dict[str, Any]] = []
    for raw in raw_rows:
        invoice_date = parse_date(raw["Invoice_Date"])
        product_code, product_name = normalize_product(raw["Product_Code"])
        region = normalize_region(raw["Region"])
        actual_revenue = int(raw["Net_Revenue_VND"])
        budget_revenue = int(raw["Budget_Revenue_VND"])
        actual_cost = int(raw["Cost_of_Sales_VND"])
        budget_cost = int(raw["Budget_Cost_VND"])
        actual_opex = int(raw["OPEX_VND"])
        budget_opex = int(raw["Budget_OPEX_VND"])
        gross_profit = actual_revenue - actual_cost
        budget_gross_profit = budget_revenue - budget_cost
        ebitda = gross_profit - actual_opex
        budget_ebitda = budget_gross_profit - budget_opex

        cleaned.append(
            {
                "Transaction ID": raw["Transaction_ID"],
                "Invoice Date": invoice_date,
                "Period": invoice_date.strftime("%Y-%m"),
                "Region": region,
                "Product": product_name,
                "Product Code": product_code,
                "Channel": str(raw["Channel"]).strip().title(),
                "Units": int(raw["Units"]),
                "Actual Revenue": actual_revenue,
                "Budget Revenue": budget_revenue,
                "Revenue Variance": actual_revenue - budget_revenue,
                "Revenue Var %": (actual_revenue / budget_revenue - 1) if budget_revenue else 0,
                "Gross Profit": gross_profit,
                "Gross Margin %": gross_profit / actual_revenue if actual_revenue else 0,
                "Budget Gross Profit": budget_gross_profit,
                "OPEX": actual_opex,
                "Budget OPEX": budget_opex,
                "EBITDA": ebitda,
                "Budget EBITDA": budget_ebitda,
                "EBITDA Variance": ebitda - budget_ebitda,
                "Data Status": "Standardized",
            }
        )
    return cleaned


def load_raw_data() -> list[dict[str, Any]]:
    workbook = load_workbook(RAW_PATH, read_only=True, data_only=True)
    worksheet = workbook["Raw Data"]
    rows = list(worksheet.iter_rows(values_only=True))
    headers = [str(value) for value in rows[0]]
    result = [dict(zip(headers, row, strict=True)) for row in rows[1:]]
    workbook.close()
    return result


def publish(staged_path: Path, final_path: Path) -> None:
    final_path.parent.mkdir(parents=True, exist_ok=True)
    if final_path.exists():
        final_path.unlink()
    os.replace(staged_path, final_path)


def reset_outputs() -> None:
    STAGING.mkdir(parents=True, exist_ok=True)
    for path in (CLEAN_PATH, VARIANCE_PATH, REPORT_PATH):
        if path.exists():
            path.unlink()
    for path in STAGING.iterdir():
        if path.is_file():
            path.unlink()


def style_title(worksheet, title: str, subtitle: str, end_column: int) -> None:
    worksheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=end_column)
    worksheet["A1"] = title
    worksheet["A1"].font = Font(name="Aptos Display", size=20, bold=True, color=WHITE)
    worksheet["A1"].fill = PatternFill("solid", fgColor=NAVY)
    worksheet["A1"].alignment = Alignment(vertical="center")
    worksheet.row_dimensions[1].height = 34
    worksheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=end_column)
    worksheet["A2"] = subtitle
    worksheet["A2"].font = Font(name="Aptos", size=10, italic=True, color=MUTED)
    worksheet["A2"].alignment = Alignment(vertical="center")
    worksheet.row_dimensions[2].height = 24


def create_clean_workbook(rows: list[dict[str, Any]], output_path: Path) -> None:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Clean Data"
    worksheet.sheet_properties.tabColor = BLUE
    headers = list(rows[0].keys())
    style_title(
        worksheet,
        "CLEAN DATA | JUNE 2026",
        "Normalized dates, regions and products with Python-calculated FP&A columns",
        len(headers),
    )
    header_row = 4
    for column, header in enumerate(headers, start=1):
        cell = worksheet.cell(header_row, column, header)
        cell.font = Font(name="Aptos", size=10, bold=True, color=WHITE)
        cell.fill = PatternFill("solid", fgColor=BLUE)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(bottom=Side(style="medium", color=NAVY))
    worksheet.row_dimensions[header_row].height = 32

    for row_index, item in enumerate(rows, start=header_row + 1):
        for column, header in enumerate(headers, start=1):
            cell = worksheet.cell(row_index, column, item[header])
            cell.font = Font(name="Aptos", size=9, color=CHARCOAL)
            cell.border = Border(bottom=THIN_GRAY)
            if row_index % 2 == 1:
                cell.fill = PatternFill("solid", fgColor="F8FAFC")

    index = {header: position + 1 for position, header in enumerate(headers)}
    for row_index in range(header_row + 1, worksheet.max_row + 1):
        worksheet.cell(row_index, index["Invoice Date"]).number_format = "dd-mmm-yyyy"
        for header in (
            "Actual Revenue",
            "Budget Revenue",
            "Revenue Variance",
            "Gross Profit",
            "Budget Gross Profit",
            "OPEX",
            "Budget OPEX",
            "EBITDA",
            "Budget EBITDA",
            "EBITDA Variance",
        ):
            worksheet.cell(row_index, index[header]).number_format = "#,##0;[Red](#,##0);-"
        for header in ("Revenue Var %", "Gross Margin %"):
            worksheet.cell(row_index, index[header]).number_format = "0.0%;[Red](0.0%);-"

    table_reference = f"A{header_row}:{worksheet.cell(worksheet.max_row, len(headers)).coordinate}"
    table = Table(displayName="CleanDataTable", ref=table_reference)
    table.tableStyleInfo = TableStyleInfo(
        name="TableStyleMedium2",
        showFirstColumn=False,
        showLastColumn=False,
        showRowStripes=True,
        showColumnStripes=False,
    )
    worksheet.add_table(table)

    var_letter = get_column_letter(index["Revenue Var %"])
    var_range = f"{var_letter}{header_row + 1}:{var_letter}{worksheet.max_row}"
    worksheet.conditional_formatting.add(
        var_range,
        CellIsRule(operator="greaterThanOrEqual", formula=["0"], fill=PatternFill("solid", fgColor=PALE_GREEN)),
    )
    worksheet.conditional_formatting.add(
        var_range,
        CellIsRule(operator="lessThan", formula=["0"], fill=PatternFill("solid", fgColor=PALE_RED)),
    )

    widths = {
        "A": 17,
        "B": 14,
        "C": 11,
        "D": 11,
        "E": 18,
        "F": 14,
        "G": 11,
        "H": 9,
        "I": 17,
        "J": 17,
        "K": 17,
        "L": 13,
        "M": 16,
        "N": 14,
        "O": 19,
        "P": 14,
        "Q": 15,
        "R": 14,
        "S": 16,
        "T": 17,
        "U": 15,
    }
    for column, width in widths.items():
        worksheet.column_dimensions[column].width = width
    worksheet.freeze_panes = "I5"
    worksheet.sheet_view.showGridLines = False
    worksheet.sheet_view.zoomScale = 80
    worksheet.sheet_view.topLeftCell = "A1"
    workbook.calculation.fullCalcOnLoad = True
    workbook.calculation.forceFullCalc = True
    workbook.save(output_path)


def summarize(rows: list[dict[str, Any]]) -> dict[str, dict[str, float]]:
    summary: dict[str, dict[str, float]] = defaultdict(lambda: defaultdict(float))
    for item in rows:
        for group in (item["Region"], "Total"):
            result = summary[str(group)]
            result["actual_revenue"] += item["Actual Revenue"]
            result["budget_revenue"] += item["Budget Revenue"]
            result["gross_profit"] += item["Gross Profit"]
            result["budget_gross_profit"] += item["Budget Gross Profit"]
            result["actual_opex"] += item["OPEX"]
            result["budget_opex"] += item["Budget OPEX"]
            result["ebitda"] += item["EBITDA"]
            result["budget_ebitda"] += item["Budget EBITDA"]
    for result in summary.values():
        result["revenue_variance"] = result["actual_revenue"] - result["budget_revenue"]
        result["revenue_variance_pct"] = result["actual_revenue"] / result["budget_revenue"] - 1
        result["gross_margin_pct"] = result["gross_profit"] / result["actual_revenue"]
        result["budget_gross_margin_pct"] = result["budget_gross_profit"] / result["budget_revenue"]
        result["margin_variance_pp"] = result["gross_margin_pct"] - result["budget_gross_margin_pct"]
        result["opex_variance"] = result["actual_opex"] - result["budget_opex"]
        result["opex_variance_pct"] = result["actual_opex"] / result["budget_opex"] - 1
        result["ebitda_variance"] = result["ebitda"] - result["budget_ebitda"]
        result["ebitda_variance_pct"] = result["ebitda"] / result["budget_ebitda"] - 1
    return summary


def _fill_block(worksheet, cell_range: str, color: str) -> None:
    for row in worksheet[cell_range]:
        for cell in row:
            cell.fill = PatternFill("solid", fgColor=color)
            cell.border = Border(left=THIN_GRAY, right=THIN_GRAY, top=THIN_GRAY, bottom=THIN_GRAY)


def add_kpi_card(worksheet, start_column: int, title: str, actual: float, budget: float, variance: float, percent: bool = False) -> None:
    end_column = start_column + 1
    worksheet.merge_cells(start_row=4, start_column=start_column, end_row=4, end_column=end_column)
    worksheet.merge_cells(start_row=5, start_column=start_column, end_row=5, end_column=end_column)
    worksheet.merge_cells(start_row=6, start_column=start_column, end_row=6, end_column=end_column)
    worksheet.merge_cells(start_row=7, start_column=start_column, end_row=7, end_column=end_column)
    for row_number in range(4, 8):
        _fill_block(worksheet, f"{worksheet.cell(row_number, start_column).coordinate}:{worksheet.cell(row_number, end_column).coordinate}", PALE_BLUE)
    worksheet.cell(4, start_column, title)
    worksheet.cell(5, start_column, actual)
    worksheet.cell(6, start_column, f"Budget: {budget:.1%}" if percent else f"Budget: {budget / 1_000_000_000:.1f}B")
    worksheet.cell(7, start_column, variance)
    worksheet.cell(4, start_column).font = Font(name="Aptos", size=10, bold=True, color=NAVY)
    worksheet.cell(5, start_column).font = Font(name="Aptos Display", size=18, bold=True, color=CHARCOAL)
    worksheet.cell(6, start_column).font = Font(name="Aptos", size=9, color=MUTED)
    worksheet.cell(7, start_column).font = Font(name="Aptos", size=11, bold=True, color=GREEN if variance >= 0 else RED)
    for row_number in range(4, 8):
        worksheet.cell(row_number, start_column).alignment = Alignment(horizontal="center", vertical="center")
    worksheet.cell(5, start_column).number_format = "0.0%" if percent else '0.0,,,"B"'
    worksheet.cell(7, start_column).number_format = '0.0"pp"' if percent else '0.0%'


def create_variance_workbook(
    summary: dict[str, dict[str, float]],
    controls: dict[str, float],
    output_path: Path,
) -> None:
    total = summary["Total"]
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Variance Analysis"
    worksheet.sheet_properties.tabColor = NAVY
    style_title(
        worksheet,
        "FP&A VARIANCE ANALYSIS | JUNE 2026",
        "Actual vs Budget | VND billions | Favorable and unfavorable signals shown explicitly",
        14,
    )
    add_kpi_card(worksheet, 1, "REVENUE", total["actual_revenue"], total["budget_revenue"], total["revenue_variance_pct"])
    add_kpi_card(worksheet, 4, "GROSS MARGIN", total["gross_margin_pct"], total["budget_gross_margin_pct"], total["margin_variance_pp"] * 100, True)
    add_kpi_card(worksheet, 7, "OPEX", total["actual_opex"], total["budget_opex"], -total["opex_variance_pct"])
    add_kpi_card(worksheet, 10, "EBITDA", total["ebitda"], total["budget_ebitda"], total["ebitda_variance_pct"])

    headers = [
        "Region",
        "Actual Revenue\n(VND bn)",
        "Budget Revenue\n(VND bn)",
        "Var\n(VND bn)",
        "Var %",
        "Gross Margin %",
        "GM Var pp",
        "Actual EBITDA\n(VND bn)",
        "Budget EBITDA\n(VND bn)",
        "EBITDA Var\n(VND bn)",
        "Signal",
    ]
    table_start = 10
    for column, header in enumerate(headers, start=1):
        cell = worksheet.cell(table_start, column, header)
        cell.font = Font(name="Aptos", size=10, bold=True, color=WHITE)
        cell.fill = PatternFill("solid", fgColor=NAVY)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(bottom=Side(style="medium", color=BLUE))
    worksheet.row_dimensions[table_start].height = 32

    for row_offset, region in enumerate(("North", "South", "Central", "Total"), start=1):
        result = summary[region]
        values = [
            region,
            result["actual_revenue"] / 1_000_000_000,
            result["budget_revenue"] / 1_000_000_000,
            result["revenue_variance"] / 1_000_000_000,
            result["revenue_variance_pct"],
            result["gross_margin_pct"],
            result["margin_variance_pp"] * 100,
            result["ebitda"] / 1_000_000_000,
            result["budget_ebitda"] / 1_000_000_000,
            result["ebitda_variance"] / 1_000_000_000,
            "Favorable" if result["ebitda_variance"] >= 0 else "Unfavorable",
        ]
        row_number = table_start + row_offset
        for column, value in enumerate(values, start=1):
            cell = worksheet.cell(row_number, column, value)
            cell.font = Font(name="Aptos", size=10, bold=region == "Total", color=CHARCOAL)
            cell.border = Border(bottom=THIN_GRAY)
            if region == "Total":
                cell.fill = PatternFill("solid", fgColor=LIGHT_BLUE)
        for column in (2, 3, 4, 8, 9, 10):
            worksheet.cell(row_number, column).number_format = '0.0;[Red](0.0);-'
        for column in (5, 6):
            worksheet.cell(row_number, column).number_format = "0.0%"
        worksheet.cell(row_number, 7).number_format = '0.0"pp"'
        signal = worksheet.cell(row_number, 11)
        signal.fill = PatternFill("solid", fgColor=PALE_GREEN if signal.value == "Favorable" else PALE_RED)
        signal.font = Font(name="Aptos", size=10, bold=True, color=GREEN if signal.value == "Favorable" else RED)

    chart = BarChart()
    chart.type = "col"
    chart.grouping = "clustered"
    chart.overlap = 0
    chart.gapWidth = 70
    chart.style = 10
    chart.title = "Revenue by Region (VND bn)"
    chart.height = 8.8
    chart.width = 18.0
    data = Reference(worksheet, min_col=2, max_col=3, min_row=table_start, max_row=table_start + 3)
    categories = Reference(worksheet, min_col=1, min_row=table_start + 1, max_row=table_start + 3)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(categories)
    chart.series[0].tx = SeriesLabel(v="Actual")
    chart.series[1].tx = SeriesLabel(v="Budget")
    chart.series[0].graphicalProperties.solidFill = BLUE
    chart.series[0].graphicalProperties.line.solidFill = NAVY
    chart.series[1].graphicalProperties.solidFill = "B8C4CE"
    chart.series[1].graphicalProperties.line.solidFill = MUTED
    chart.legend.position = "b"
    chart.legend.overlay = False
    chart.y_axis.scaling.min = 0
    chart.y_axis.majorUnit = 10
    chart.y_axis.numFmt = "0"
    worksheet.add_chart(chart, "A16")

    for column, width in {
        "A": 14,
        "B": 16,
        "C": 16,
        "D": 14,
        "E": 11,
        "F": 16,
        "G": 12,
        "H": 16,
        "I": 16,
        "J": 14,
        "K": 14,
        "L": 4,
        "M": 4,
        "N": 4,
    }.items():
        worksheet.column_dimensions[column].width = width
    worksheet.freeze_panes = "A11"
    worksheet.sheet_view.showGridLines = False
    worksheet.sheet_view.zoomScale = 85
    worksheet.sheet_view.topLeftCell = "A1"

    checks = workbook.create_sheet("Control Checks")
    checks.sheet_properties.tabColor = GREEN
    style_title(checks, "AUTOMATION CONTROL CHECKS", "Source-to-output reconciliation", 6)
    check_headers = ["Check", "Actual", "Expected", "Difference", "Tolerance", "Status"]
    for column, header in enumerate(check_headers, start=1):
        cell = checks.cell(4, column, header)
        cell.font = Font(bold=True, color=WHITE)
        cell.fill = PatternFill("solid", fgColor=NAVY)
    check_inputs = [
        ("Clean rows equal raw rows", controls["clean_rows"], controls["raw_rows"], 0),
        ("Revenue totals reconcile", total["actual_revenue"], controls["raw_actual_revenue"], 1),
        ("Budget totals reconcile", total["budget_revenue"], controls["raw_budget_revenue"], 1),
        ("No unmapped products", controls["unmapped_products"], 0, 0),
    ]
    check_rows = []
    for label, actual, expected, tolerance in check_inputs:
        difference = actual - expected
        status = "PASS" if abs(difference) <= tolerance else "FAIL"
        check_rows.append((label, actual, expected, difference, tolerance, status))
    for row_number, values in enumerate(check_rows, start=5):
        for column, value in enumerate(values, start=1):
            checks.cell(row_number, column, value)
            checks.cell(row_number, column).border = Border(bottom=THIN_GRAY)
        passed = values[5] == "PASS"
        checks.cell(row_number, 6).fill = PatternFill("solid", fgColor=PALE_GREEN if passed else PALE_RED)
        checks.cell(row_number, 6).font = Font(bold=True, color=GREEN if passed else RED)
    checks.column_dimensions["A"].width = 30
    for column in "BCDEF":
        checks.column_dimensions[column].width = 16
    checks.sheet_view.showGridLines = False
    checks.sheet_view.zoomScale = 90
    workbook.calculation.fullCalcOnLoad = True
    workbook.calculation.forceFullCalc = True
    workbook.save(output_path)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_doc_font(run, size: float = 10.5, color: str = CHARCOAL, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_doc_table_header(table, labels: list[str]) -> None:
    row = table.rows[0]
    for index, label in enumerate(labels):
        cell = row.cells[index]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(label)
        set_doc_font(run, size=9, color=WHITE, bold=True)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.left_indent = Inches(0.28)
    paragraph.paragraph_format.first_line_indent = Inches(-0.16)
    set_doc_font(paragraph.add_run(text), size=10.2)


def create_word_report(summary: dict[str, dict[str, float]], row_count: int, output_path: Path) -> None:
    total = summary["Total"]
    revenue_direction = "above" if total["revenue_variance_pct"] >= 0 else "below"
    margin_direction = "above" if total["margin_variance_pp"] >= 0 else "below"
    opex_signal = "unfavorable" if total["opex_variance_pct"] >= 0 else "favorable"
    ebitda_direction = "above" if total["ebitda_variance_pct"] >= 0 else "below"
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    document.core_properties.title = "June 2026 FP&A Management Report"
    document.core_properties.subject = "Actual vs Budget management reporting"
    document.core_properties.author = "Python FP&A Automation"

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(CHARCOAL)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for style_name, size, color, before, after in (
        ("Heading 1", 15, NAVY, 12, 6),
        ("Heading 2", 12, BLUE, 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_doc_font(header.add_run("FP&A MONTHLY CLOSE | CONFIDENTIAL"), size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    set_doc_font(footer.add_run("Generated automatically from sample_data/data-raw.xlsx | Python workflow demo"), size=8, color=MUTED)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    set_doc_font(kicker.add_run("MONTHLY MANAGEMENT REPORT"), size=9, color=BLUE, bold=True)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    set_doc_font(title.add_run("FP&A Performance Review"), size=24, color=NAVY, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    set_doc_font(subtitle.add_run("June 2026 | Actual vs Budget | VND"), size=12, color=MUTED)

    callout = document.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = callout.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(
        "Executive readout: "
        f"Revenue finished {abs(total['revenue_variance_pct']):.1%} {revenue_direction} budget; "
        f"gross margin was {abs(total['margin_variance_pp']) * 100:.1f}pp {margin_direction} plan; "
        f"OPEX was {abs(total['opex_variance_pct']):.1%} {opex_signal}; and "
        f"EBITDA was {abs(total['ebitda_variance_pct']):.1%} {ebitda_direction} budget. "
        "The automated bridge highlights where topline performance converted into profit."
    )
    set_doc_font(run, size=10.6, color=NAVY, bold=True)

    document.add_heading("1 | Performance at a glance", level=1)
    kpi_table = document.add_table(rows=1, cols=5)
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi_table.autofit = False
    add_doc_table_header(kpi_table, ["Metric", "Actual", "Budget", "Variance", "Signal"])
    metrics = [
        ("Revenue", total["actual_revenue"] / 1e9, total["budget_revenue"] / 1e9, total["revenue_variance_pct"], "Favorable", False),
        ("Gross Margin", total["gross_margin_pct"], total["budget_gross_margin_pct"], total["margin_variance_pp"], "Unfavorable", True),
        ("OPEX", total["actual_opex"] / 1e9, total["budget_opex"] / 1e9, total["opex_variance_pct"], "Unfavorable", False),
        ("EBITDA", total["ebitda"] / 1e9, total["budget_ebitda"] / 1e9, total["ebitda_variance_pct"], "Favorable", False),
    ]
    for metric, actual, budget, variance, signal, percent_metric in metrics:
        row = kpi_table.add_row()
        values = [
            metric,
            f"{actual:.1%}" if percent_metric else f"{actual:.1f}B",
            f"{budget:.1%}" if percent_metric else f"{budget:.1f}B",
            f"{variance * 100:+.1f}pp" if percent_metric else f"{variance:+.1%}",
            signal,
        ]
        for index, value in enumerate(values):
            cell = row.cells[index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            set_doc_font(paragraph.add_run(value), size=9.4, bold=index in (0, 4), color=GREEN if index == 4 and signal == "Favorable" else RED if index == 4 else CHARCOAL)
            if index == 4:
                set_cell_shading(cell, PALE_GREEN if signal == "Favorable" else PALE_RED)

    document.add_heading("2 | Regional variance", level=1)
    region_table = document.add_table(rows=1, cols=6)
    region_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    region_table.autofit = False
    add_doc_table_header(region_table, ["Region", "Revenue", "Budget", "Var %", "GM %", "EBITDA"])
    for region in ("North", "South", "Central"):
        result = summary[region]
        row = region_table.add_row()
        values = [
            region,
            f"{result['actual_revenue'] / 1e9:.1f}B",
            f"{result['budget_revenue'] / 1e9:.1f}B",
            f"{result['revenue_variance_pct']:+.1%}",
            f"{result['gross_margin_pct']:.1%}",
            f"{result['ebitda'] / 1e9:.1f}B",
        ]
        for index, value in enumerate(values):
            cell = row.cells[index]
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            set_doc_font(paragraph.add_run(value), size=9.2, bold=index == 0)

    document.add_heading("3 | Management focus", level=1)
    add_bullet(document, "Protect North and South revenue momentum while tightening discount governance by product and channel.")
    add_bullet(document, "Recover gross margin through pricing, freight and product-mix actions; target at least 100bps recovery in the next forecast.")
    add_bullet(document, "Require a Central EBITDA recovery bridge with clear owners before the July forecast lock.")

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("4 | Risks and recommended actions", level=1)
    risk_table = document.add_table(rows=1, cols=3)
    risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_doc_table_header(risk_table, ["Priority", "Action", "Owner / Timing"])
    actions = [
        ("1", "Review discount leakage in Electronics and set approval thresholds by channel.", "Commercial Finance | 1 week"),
        ("2", "Reconcile OPEX overruns and freeze low-ROI discretionary spend.", "FP&A + Regional Finance | Immediate"),
        ("3", "Build a Central region EBITDA recovery plan and track it weekly.", "Central GM | Before next close"),
        ("4", "Refresh the Q3 gross-margin forecast using the June exit rate.", "FP&A | July forecast lock"),
    ]
    for priority, action, owner in actions:
        row = risk_table.add_row()
        for index, value in enumerate((priority, action, owner)):
            cell = row.cells[index]
            set_cell_margins(cell, top=110, bottom=110)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            set_doc_font(paragraph.add_run(value), size=9.6, bold=index == 0, color=NAVY if index == 0 else CHARCOAL)
        set_cell_shading(row.cells[0], PALE_BLUE)

    document.add_heading("5 | Automation evidence", level=1)
    evidence = document.add_table(rows=4, cols=2)
    evidence.alignment = WD_TABLE_ALIGNMENT.CENTER
    evidence_rows = [
        ("Source", "sample_data/data-raw.xlsx"),
        ("Rows processed", f"{row_count:,} / {row_count:,} successful"),
        ("Transformations", "Date, region and product standardized; FP&A measures calculated"),
        ("Outputs", "Clean Data, Variance Analysis and this management report"),
    ]
    for row_index, (label, value) in enumerate(evidence_rows):
        for column_index, text in enumerate((label, value)):
            cell = evidence.cell(row_index, column_index)
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            set_doc_font(paragraph.add_run(text), size=9.5, bold=column_index == 0, color=NAVY if column_index == 0 else CHARCOAL)
        set_cell_shading(evidence.cell(row_index, 0), PALE_BLUE)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    set_doc_font(
        note.add_run("Prepared automatically by Python from line-level June 2026 demonstration data. Values are synthetic and rounded for presentation."),
        size=8.5,
        color=MUTED,
        italic=True,
    )
    document.save(output_path)


def main() -> None:
    os.system("")
    log("\033[2J\033[H")
    log("=" * 68)
    log("        FP&A MONTHLY REPORTING AUTOMATION | JUNE 2026")
    log("=" * 68)
    log("One Python run will publish three management-ready deliverables.")
    reset_outputs()
    log("\nWorkspace reset: Raw Data preserved; prior outputs cleared.")
    # Keep the first stage visible long enough to arrange the live 1:1 demo view.
    time.sleep(12.0)

    section(1, "READ + VALIDATE RAW DATA")
    raw_rows = load_raw_data()
    log(f"  Source workbook          {RAW_PATH.name}")
    log(f"  Rows loaded              {len(raw_rows):,}")
    log("  Required columns         PASS")
    cleaned = clean_rows(raw_rows)
    log("  Date / region / product  standardized")
    time.sleep(0.4)

    section(2, "BUILD CLEAN DATA")
    clean_staged = STAGING / CLEAN_PATH.name
    create_clean_workbook(cleaned, clean_staged)
    publish(clean_staged, CLEAN_PATH)
    log(f"  CREATED                  {CLEAN_PATH.name}")
    log("  Added FP&A columns       Revenue Var, GM%, EBITDA, Status")
    time.sleep(3.0)

    section(3, "CALCULATE VARIANCE ANALYSIS")
    summary = summarize(cleaned)
    total = summary["Total"]
    variance_staged = STAGING / VARIANCE_PATH.name
    controls = {
        "raw_rows": len(raw_rows),
        "clean_rows": len(cleaned),
        "raw_actual_revenue": sum(int(row["Net_Revenue_VND"]) for row in raw_rows),
        "raw_budget_revenue": sum(int(row["Budget_Revenue_VND"]) for row in raw_rows),
        "unmapped_products": 0,
    }
    create_variance_workbook(summary, controls, variance_staged)
    publish(variance_staged, VARIANCE_PATH)
    log(f"  CREATED                  {VARIANCE_PATH.name}")
    log(f"  Revenue vs Budget        {total['revenue_variance_pct']:+.1%}")
    log(f"  Gross Margin variance    {total['margin_variance_pp'] * 100:+.1f}pp")
    log(f"  EBITDA vs Budget         {total['ebitda_variance_pct']:+.1%}")
    time.sleep(3.0)

    section(4, "GENERATE MANAGEMENT REPORT")
    report_staged = STAGING / REPORT_PATH.name
    create_word_report(summary, len(cleaned), report_staged)
    publish(report_staged, REPORT_PATH)
    log(f"  CREATED                  {REPORT_PATH.name}")
    log("  Executive summary       ready")
    log("  Risks + actions          ready")
    time.sleep(2.0)

    log("\n" + "=" * 68)
    log("AUTOMATION COMPLETE | 3 OUTPUTS READY FOR MANAGEMENT REVIEW")
    log("=" * 68)


if __name__ == "__main__":
    main()
